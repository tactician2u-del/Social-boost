import streamlit as st
from textblob import TextBlob
import pandas as pd
import plotly.express as px
from datetime import datetime
import io
from docx import Document

st.set_page_config(page_title="SocialBoost AI Chatbot", layout="wide")
st.title("🚀 SocialBoost AI — Chatbot for Social Media Campaigns")
st.markdown("**Reach • Engagement • Positive Sentiment Booster**")

# Sidebar — Campaign Setup
with st.sidebar:
    st.header("📋 Campaign Setup")
    campaign_name = st.text_input("Campaign Name", "Election Sentiment Drive 2026")
    platforms = st.multiselect("Platforms", ["X (Twitter)", "Instagram", "LinkedIn", "Facebook", "TikTok"], default=["X (Twitter)", "Instagram"])
    audience = st.text_input("Target Audience", "Young voters, urban & rural")
    goal = st.selectbox("Main Goal", ["Increase positive sentiment", "Maximize reach", "Boost engagement", "Drive shares/comments"])
    st.caption("AI will tailor every response to these settings")

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": f"Hi! I'm SocialBoost AI for **{campaign_name}**. How can I help boost your reach, engagement, and positive sentiment today? 🎯"}
    ]

# Display chat history
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Quick action buttons
col1, col2, col3 = st.columns(3)
if col1.button("📝 Generate 5 Sample Posts"):
    st.session_state.messages.append({"role": "user", "content": "Generate 5 sample posts"})
if col2.button("📈 Analyze Past Posts CSV"):
    st.session_state.show_analyzer = True
if col3.button("💡 Get Engagement Tips"):
    st.session_state.messages.append({"role": "user", "content": "Give me top tips to increase reach and positive sentiment"})

# Chat input
if prompt := st.chat_input("Ask anything about your campaign..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # Generate AI response
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            response = generate_response(prompt, platforms, audience, goal)
            st.markdown(response)
    st.session_state.messages.append({"role": "assistant", "content": response})

# Helper function
def generate_response(prompt, platforms, audience, goal):
    prompt_lower = prompt.lower()
    if "generate" in prompt_lower or "post" in prompt_lower or "caption" in prompt_lower:
        return f"Here are 5 optimized {platforms[0]} posts for **{audience}** aimed at **{goal}**:\n\n1. 'What do YOU think about the latest election survey? Share your voice! 🇺🇸 #PositiveVote #EngageNow'\n2. ... (full 5 posts with emojis, hashtags, CTAs — tailored for positive sentiment)"
    elif any(word in prompt_lower for word in ["analyze", "sentiment", "score"]):
        # Demo: if user pasted text, analyze it
        sample_text = prompt.split("analyze")[-1].strip() if "analyze" in prompt_lower else "Sample positive campaign post here"
        blob = TextBlob(sample_text)
        polarity = blob.sentiment.polarity
        sentiment = "😊 Highly Positive" if polarity > 0.3 else "😐 Neutral" if polarity > -0.1 else "😞 Negative"
        rewrite = "Rewritten for max positivity: " + sample_text.replace("bad", "amazing")  # simple demo rewrite
        return f"**Sentiment Analysis**: {sentiment} (score: {polarity:.2f})\n**Suggestion to boost positivity**: {rewrite}"
    elif "engagement" in prompt_lower or "reach" in prompt_lower or "tip" in prompt_lower:
        return f"**Top 5 tips for {goal} on {', '.join(platforms)}**:\n• Post at peak times (8–10 AM & 6–8 PM)\n• Use 3–5 relevant hashtags\n• Add questions & emojis for 3x comments\n• A/B test 2 versions\n• Reply to every comment within 1 hour"
    else:
        return f"Got it! For your **{campaign_name}** targeting **{audience}**, here's how we can drive **{goal}**: [smart, personalized strategy based on your input]"

# Past Posts Analyzer Tab (optional full-screen mode)
if st.session_state.get("show_analyzer", False):
    st.subheader("📊 Past Posts Sentiment Analyzer")
    uploaded = st.file_uploader("Upload CSV of past posts (columns: Post_Text, Date, Platform)", type=["csv"])
    if uploaded:
        df = pd.read_csv(uploaded)
        df["Sentiment"] = df["Post_Text"].apply(lambda x: TextBlob(str(x)).sentiment.polarity)
        st.dataframe(df)
        fig = px.histogram(df, x="Sentiment", color="Platform", title="Sentiment Distribution")
        st.plotly_chart(fig)
        st.success("Upload any number of CSVs — analyze as many campaigns as you like!")

# Exports
st.divider()
st.subheader("📤 Export Campaign")
col_exp1, col_exp2 = st.columns(2)
with col_exp1:
    if st.button("📘 Download Full Chat + Strategy as Word"):
        doc = Document()
        doc.add_heading(campaign_name, 0)
        for msg in st.session_state.messages:
            doc.add_paragraph(f"{msg['role'].upper()}: {msg['content']}")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("⬇️ Download Word", buffer, f"{campaign_name}_strategy.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
with col_exp2:
    if st.button("📕 Download as PDF"):
        # Simple PDF via reportlab (add to requirements if you want full charts)
        st.info("PDF export ready in next update — for now use browser Print → Save as PDF")

st.caption("Refresh page anytime for unlimited new sessions • No data stored")
